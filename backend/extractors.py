import httpx
from bs4 import BeautifulSoup
from readability import Document
import docx
import pypdf
import io
from typing import Optional
from fastapi import UploadFile

async def extract_text_from_url(url: str) -> str:
    """Extract main content from URL"""
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(url, timeout=30.0)
            response.raise_for_status()
            
        # Use readability to extract main content
        doc = Document(response.text)
        soup = BeautifulSoup(doc.summary(), 'html.parser')
        
        # Clean up the text
        text = soup.get_text()
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n'.join(lines)
        
    except Exception as e:
        return f"Error extracting from URL: {str(e)}"

async def extract_text_from_docx(file: UploadFile) -> str:
    """Extract text from DOCX file"""
    try:
        content = await file.read()
        doc = docx.Document(io.BytesIO(content))
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        return '\n'.join(text_parts)
        
    except Exception as e:
        return f"Error extracting from DOCX: {str(e)}"

async def extract_text_from_pdf(file: UploadFile) -> str:
    """Extract text from PDF file"""
    try:
        content = await file.read()
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
        
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text.strip())
        
        return '\n'.join(text_parts)
        
    except Exception as e:
        return f"Error extracting from PDF: {str(e)}"
