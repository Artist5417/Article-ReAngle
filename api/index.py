from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from typing import Optional
import requests
import os
import io
import httpx
from bs4 import BeautifulSoup
from readability import Document
import docx
import pypdf

# 创建FastAPI应用
app = FastAPI(title='Article ReAngle')

# 添加CORS中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=['*'],
    allow_credentials=True,
    allow_methods=['*'],
    allow_headers=['*'],
)

async def extract_text_from_url(url: str) -> str:
    """Extract main content from URL"""
    try:
        print(f"开始提取URL内容: {url}")
        
        # 确保URL有协议
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
            print(f"添加协议后的URL: {url}")
        
        async with httpx.AsyncClient() as client:
            print("发送HTTP请求...")
            response = await client.get(url, timeout=30.0)
            print(f"HTTP响应状态: {response.status_code}")
            response.raise_for_status()
        
        print("开始解析HTML内容...")
        doc = Document(response.text)
        summary = doc.summary()
        print(f"Readability提取的摘要长度: {len(summary)}")
        
        soup = BeautifulSoup(summary, 'html.parser')
        text = soup.get_text()
        
        # 清理文本
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        result = '\n'.join(lines)
        print(f"最终提取的文本长度: {len(result)}")
        
        return result
        
    except Exception as e:
        print(f"URL提取错误: {str(e)}")
        return f"Error extracting from URL: {str(e)}"

async def extract_text_from_docx(file: UploadFile) -> str:
    """Extract text from DOCX file"""
    try:
        content = await file.read()
        doc = docx.Document(io.BytesIO(content))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        return f"Error extracting from DOCX: {str(e)}"

async def extract_text_from_pdf(file: UploadFile) -> str:
    """Extract text from PDF file"""
    try:
        content = await file.read()
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)
    except Exception as e:
        return f"Error extracting from PDF: {str(e)}"

def rewrite_text_with_openai(text: str, prompt: str, api_key: str) -> str:
    """使用OpenAI API重写文本"""
    try:
        if not api_key:
            return "错误：未提供OpenAI API Key"
        
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
        
        # 根据原文长度动态设置目标与上限（先按大模型上限计算）
        text_length = len(text)
        # 目标生成长度：原文约110%，以保证不变短；保底1000
        target_tokens = max(1000, int(text_length * 1.1))
        # 固定使用更高上限的模型
        model = 'gpt-4o-mini'
        # 经验上限：gpt-4o-mini 总窗口约 20000，生成上限≈16000
        model_total_window = 20000
        model_completion_cap = 16000
        # 估算输入token（粗略按4字符≈1token），并为系统/提示保留额外预算
        system_and_overhead_tokens = 300
        approx_input_tokens = int(len(text) / 4) + system_and_overhead_tokens

        # 可用生成空间 = 总窗口 - 输入 - 安全余量
        safety_margin = 500
        available_tokens = max(0, model_total_window - approx_input_tokens - safety_margin)

        # 生成上限受三个因素共同约束：目标、模型生成上限、可用空间
        max_tokens = max(
            256,
            min(target_tokens, model_completion_cap, available_tokens)
        )
        
        # 构造统一提示词（忠实改写器）
        system_prompt = (
            "你是一个“忠实改写器”，也是一个精准且不编造内容的文本改写引擎。\n\n"
            "【任务】\n"
            "在不新增事实、不删除关键信息的前提下，根据用户给定的“观点/立场与表达风格”对原文进行重写。\n\n"
            "【核心原则】\n"
            "1. 内容忠实：不得编造、夸大或引入原文没有的事实、数据、引用、时间线、人名或机构名。\n"
            "2. 信息保全：保留原文的核心要点、关键数据、结论与限定条件；如原文有不确定性或前提，必须保留。\n"
            "3. 立场与风格：严格按用户指定的观点与表达方式改写（如学术、口语、营销、中立、正向、反向），通过措辞和论证角度体现立场，而不是篡改事实。\n"
            "4. 风格控制：在语气、词汇、逻辑密度、结构上体现用户要求的风格。\n"
            "5. 可读性：结构清晰、层次分明，必要时可用小标题和段落组织内容。\n"
            "6. 引文与链接：若原文包含引用、链接、人名、机构名或数值，必须完整保留且不得歧义重写。\n"
            "7. 长度控制：若用户指定字数或篇幅，需严格遵守（允许上下浮动 10%）；未指定时保持与原文相近。\n"
            "8. 语言保持：默认保持与原文相同的语言，除非用户另有要求。\n"
            "9. 冲突兜底：若用户需求与事实冲突，应保持事实，用语气和 framing 表达立场，不得改写事实本身。\n"
            "10. 敏感与合规：避免人身攻击、仇恨或违规内容；任何情况下不得生成违法内容。\n\n"
            "【输出规则】\n"
            "- 只输出改写后的正文，不要前言、解释或步骤说明。\n"
            "- 默认使用 Markdown 格式（保留段落、列表、引用、代码块等）。"
        )

        user_prompt_template = (
            "【用户需求】\n{req}\n\n"
            "【原文】\n{src}\n\n"
            "【任务】\n请严格遵循以上规则，根据“用户需求”重写“原文”，输出完整正文。"
        )

        data = {
            'model': model,
            'messages': [
                {
                    'role': 'system',
                    'content': system_prompt
                },
                {
                    'role': 'user',
                    'content': user_prompt_template.format(req=prompt, src=text)
                }
            ],
            'max_tokens': max_tokens,
            'temperature': 0.7
        }
        
        def call_openai(payload):
            return requests.post(
                'https://api.openai.com/v1/chat/completions',
                headers=headers,
                json=payload,
                timeout=60
            )

        # 超限前置校验：若无可用生成空间则直接报错
        if available_tokens < 256:
            return (
                f"错误：输入过长超出模型窗口。估算输入tokens约 {approx_input_tokens}，"
                f"模型总窗口 {model_total_window}，安全余量 {safety_margin}，可用生成空间 {available_tokens}。"
                f"请缩短输入或改用分段重写。"
            )

        # 发起一次请求（已按硬上限收紧）
        response = call_openai(data)
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            # 返回更清晰的硬上限提示
            err_text = response.text or ''
            if 'max_tokens' in err_text or 'maximum' in err_text:
                return (
                    "OpenAI API错误: 请求超出模型限制。" 
                    f"输入估算 {approx_input_tokens} tokens，max_tokens 请求 {max_tokens}，"
                    f"模型生成上限约 {model_completion_cap}，总窗口 {model_total_window}，"
                    f"可用生成空间 {available_tokens}。请缩短输入或采用分段重写。\n原始错误：{err_text}"
                )
            return f"OpenAI API错误: {response.status_code} - {response.text}"
            
    except Exception as e:
        return f"处理错误: {str(e)}"

@app.get('/')
async def root():
    return {"message": "Article ReAngle API is running"}

@app.get('/test')
async def test():
    return {"status": "working", "message": "Test endpoint is working"}

@app.post('/process')
async def process(
    input_text: Optional[str] = Form(None),
    url: Optional[str] = Form(None),
    prompt: Optional[str] = Form('改写成新闻报道风格'),
    api_key: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    try:
        print(f"收到处理请求 - 文本: {bool(input_text)}, URL: {bool(url)}, 文件: {bool(file)}")
        
        # 获取原始文本
        raw_text = ''
        
        if input_text:
            raw_text = input_text
            print(f"使用文本输入，长度: {len(raw_text)}")
        elif url:
            print(f"正在处理URL: {url}")
            raw_text = await extract_text_from_url(url)
            print(f"URL提取结果长度: {len(raw_text)}")
            print(f"URL提取结果前100字符: {raw_text[:100]}")
        elif file:
            print(f"正在处理文件: {file.filename}")
            if file.filename.lower().endswith('.docx'):
                raw_text = await extract_text_from_docx(file)
            elif file.filename.lower().endswith('.pdf'):
                raw_text = await extract_text_from_pdf(file)
            else:  # 处理TXT和其他文本文件
                content = await file.read()
                raw_text = content.decode('utf-8', errors='ignore')
            print(f"文件提取结果长度: {len(raw_text)}")
        
        if not raw_text:
            return JSONResponse({'error': '没有提供有效的输入内容'}, status_code=400)
        
        # 使用OpenAI重写文本
        print("开始调用OpenAI API...")
        rewritten = rewrite_text_with_openai(raw_text, prompt, api_key)
        print(f"重写完成，结果长度: {len(rewritten)}")
        
        return {
            'original': raw_text,
            'rewritten': rewritten,
            'summary': raw_text
        }
        
    except Exception as e:
        print(f"❌ 处理错误: {str(e)}")
        return JSONResponse({'error': f'处理失败: {str(e)}'}, status_code=500)