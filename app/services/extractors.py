"""
根据不同输入源提取文本
"""

import httpx
import docx
import pypdf
import io
from bs4 import BeautifulSoup
from readability import Document
from fastapi import UploadFile
from loguru import logger
import os
import tempfile
import subprocess
from pathlib import Path

from app.core.exceptions import ContentExtractionError, InvalidInputError

from PIL import Image
from urllib.parse import urlparse, parse_qs
import re

# 轻量 OCR（中英混排效果较好）
try:
    from rapidocr_onnxruntime import RapidOCR  # type: ignore
    import numpy as np  # 仅在 RapidOCR 可用时再导入 numpy
    _rapid_ocr = RapidOCR()  # 初始化一次复用
except Exception:
    _rapid_ocr = None
    np = None  # type: ignore
    logger.warning("RapidOCR is not available. Image OCR will be disabled.")

 

async def extract_text_from_url(url: str) -> str:
    """
    从 URL 提取主要文本内容。
    """
    try:
        logger.info(f"Starting URL extraction: {url}")

        # 确保URL有协议
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
            logger.debug(f"URL protocol added: {url}")

        async with httpx.AsyncClient() as client:
            logger.debug("Sending HTTP request...")
            response = await client.get(url, timeout=30.0)
            logger.debug(f"HTTP response status: {response.status_code}")
            response.raise_for_status()

        logger.debug("Parsing HTML content...")
        doc = Document(response.text)
        summary = doc.summary()
        logger.debug(f"Readability summary length: {len(summary)}")

        soup = BeautifulSoup(summary, "html.parser")
        text = soup.get_text()

        # 清理文本
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        result = "\n".join(lines)
        logger.info(f"Extraction complete. Length: {len(result)}")

        if not result:
            raise ContentExtractionError("Extracted text is empty")

        return result

    except httpx.HTTPError as e:
        logger.error(f"HTTP error during extraction: {e}")
        raise ContentExtractionError(f"Failed to download URL: {str(e)}")
    except Exception as e:
        logger.exception(f"Unexpected error extracting from URL: {url}")
        raise ContentExtractionError(f"Error extracting from URL: {str(e)}")


async def extract_text_from_docx(file: UploadFile) -> str:
    """
    从 DOCX 文件提取文本。
    """
    try:
        logger.info(f"Starting DOCX extraction: {file.filename}")
        content = await file.read()
        doc = docx.Document(io.BytesIO(content))

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        result = "\n".join(text_parts)
        logger.info(f"DOCX extraction complete. Length: {len(result)}")
        return result

    except Exception as e:
        logger.exception(f"Error extracting from DOCX: {file.filename}")
        raise ContentExtractionError(f"Error extracting from DOCX: {str(e)}")


async def extract_text_from_pdf(file: UploadFile) -> str:
    """
    从 PDF 文件提取文本。
    """
    try:
        logger.info(f"Starting PDF extraction: {file.filename}")
        content = await file.read()
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))

        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())

        result = "\n".join(text_parts)
        logger.info(f"PDF extraction complete. Length: {len(result)}")
        return result

    except Exception as e:
        logger.exception(f"Error extracting from PDF: {file.filename}")
        raise ContentExtractionError(f"Error extracting from PDF: {str(e)}")


async def extract_text_from_image(file: UploadFile) -> str:
    """
    从图片文件提取文本（OCR）。支持常见格式：png/jpg/jpeg/webp。
    优先使用 RapidOCR；若不可用则报错提示安装。
    """
    if _rapid_ocr is None:
        raise ContentExtractionError(
            "OCR not available. Please install rapidocr-onnxruntime."
        )
    try:
        logger.info(f"Starting IMAGE OCR: {file.filename}")
        content = await file.read()
        # 基础预处理：转 RGB，尽量消除模式差异
        image = Image.open(io.BytesIO(content)).convert("RGB")
        np_img = np.array(image)
        result, _ = _rapid_ocr(np_img)
        lines = []
        for item in (result or []):
            # 兼容不同版本 RapidOCR 的返回：
            # 1) [box, text, score]
            # 2) [box, (text, score)]
            text = ""
            if isinstance(item, (list, tuple)):
                if len(item) >= 3:
                    # [box, text, score]
                    text = str(item[1] or "")
                elif len(item) == 2:
                    second = item[1]
                    if isinstance(second, (list, tuple)) and second:
                        text = str(second[0] or "")
                    else:
                        text = str(second or "")
            else:
                # 非预期结构，跳过
                continue
            if text.strip():
                lines.append(text.strip())
        extracted = "\n".join(lines)
        logger.info(f"IMAGE OCR complete. Length: {len(extracted)}")
        if not extracted.strip():
            # 给到用户可读的错误，便于前端提示
            raise ContentExtractionError("OCR result is empty. Try a clearer image.")
        return extracted
    except ContentExtractionError:
        raise
    except Exception as e:
        logger.exception(f"Error extracting from IMAGE: {file.filename}")
        raise ContentExtractionError(f"Error extracting from IMAGE: {str(e)}")

# ------------------------------
# YouTube Ingestion (v1) Section
# ------------------------------
# Step 1: URL 校验与 videoId 提取（本地，不发网）
_YOUTUBE_ALLOWED_HOSTS = {"youtube.com", "www.youtube.com", "m.youtube.com", "youtu.be"}
_VIDEO_ID_RE = re.compile(r"^[A-Za-z0-9_-]{11}$")

def validate_and_get_video_id(url: str) -> str:
    """
    仅基于 URL 做本地校验并提取 11 位 videoId。
    - 只接受 http/https
    - host 必须在 YouTube 白名单
    - 支持 /watch?v=VID, youtu.be/VID, /shorts/VID, /embed/VID
    - playlist/channel/@user 等非单视频形态直接拒绝
    """
    try:
        if not url or not isinstance(url, str):
            raise InvalidInputError("YouTube URL is required")
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
        parsed = urlparse(url)
        scheme = (parsed.scheme or "").lower()
        if scheme not in ("http", "https"):
            raise InvalidInputError("Only http/https allowed for YouTube URL")
        host = (parsed.netloc or "").split("@")[-1].split(":")[0].lower()
        if host not in _YOUTUBE_ALLOWED_HOSTS:
            raise InvalidInputError("Unsupported YouTube host")

        path = parsed.path or ""
        lower_path = path.lower()
        query = parse_qs(parsed.query or "")

        # 明确拒绝 playlist / channel / user 等形态
        if "playlist" in lower_path or lower_path.startswith("/channel/") or lower_path.startswith("/c/") or lower_path.startswith("/@"):
            raise InvalidInputError("Playlist or channel URL is not supported")

        video_id = ""
        if host == "youtu.be":
            video_id = path.strip("/").split("/")[0]
        else:
            if lower_path.startswith("/watch"):
                video_id = (query.get("v") or [""])[0]
                if not video_id and "list" in query:
                    # 只有 list 而没有 v → 认为是 playlist
                    raise InvalidInputError("Playlist URL is not supported")
            elif lower_path.startswith("/shorts/") or lower_path.startswith("/embed/"):
                parts = path.split("/")
                video_id = parts[2] if len(parts) >= 3 else ""

        if not _VIDEO_ID_RE.match(video_id or ""):
            raise InvalidInputError("Invalid or missing YouTube videoId")
        return video_id
    except InvalidInputError:
        raise
    except Exception:
        raise InvalidInputError("Invalid YouTube URL")

# Step 2: 轻量探测（不下载媒体），验证可解析性并返回基础元信息
# 说明：不强依赖 Data API；使用 yt-dlp 的 extract_info(download=False)
try:
    from yt_dlp import YoutubeDL  # type: ignore
    _yt_dlp_available = True
except Exception:
    _yt_dlp_available = False
    logger.warning("yt-dlp is not available. YouTube metadata probing disabled.")

from typing import Dict, Any, List, Optional

def probe_youtube_basic_info(url: str) -> Dict[str, Any]:
    """
    使用 yt-dlp 做不下载媒体的元信息探测，判断是否为可解析的单视频：
    - 拒绝 playlist/entries
    - 拒绝直播与即将直播
    - 拒绝会员/私有/需登录/受限
    - 校验时长存在且 > 0
    成功返回:
      { videoId, title, duration, is_live, availability, age_limit }
    失败抛 ContentExtractionError。
    """
    if not _yt_dlp_available:
        raise ContentExtractionError(
            "yt-dlp not installed. Please `pip install yt-dlp` to enable YouTube probing."
        )
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    ydl_opts = {
        "quiet": True,
        "skip_download": True,
        "noplaylist": True,
        "extract_flat": "discard_in_playlist",
        "socket_timeout": 10,
    }
    try:
        with YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
        # playlist 识别
        if info.get("_type") == "playlist" or "entries" in info:
            raise ContentExtractionError("Playlist URL is not supported")
        video_id = info.get("id") or ""
        title = info.get("title") or ""
        duration = info.get("duration")
        is_live = bool(info.get("is_live") or (info.get("live_status") in {"is_live", "is_upcoming"}))
        availability = (info.get("availability") or "").lower()
        age_limit = int(info.get("age_limit") or 0)
        # 基本可播性判断
        if is_live:
            raise ContentExtractionError("Live stream is not supported")
        if not isinstance(duration, (int, float)) or int(duration) <= 0:
            raise ContentExtractionError("Video is not playable or duration is unknown")
        if availability and availability not in {"public", "unlisted"}:
            # 常见：private, needs_auth, premium_only, subscriber_only 等
            raise ContentExtractionError("Video requires membership or is restricted")
        # v1 直接拒绝年龄限制
        if age_limit >= 18:
            raise ContentExtractionError("Age restricted video is not supported in v1")
        return {
            "videoId": video_id,
            "title": title,
            "duration": int(duration),
            "is_live": is_live,
            "availability": availability,
            "age_limit": age_limit,
        }
    except ContentExtractionError:
        raise
    except Exception as e:
        logger.exception(f"yt-dlp probing failed for url={url} | reason={e}")
        raise ContentExtractionError(f"Failed to probe YouTube metadata: {str(e)}")

# Step 3: 字幕提取（仅字幕，不做 ASR）
import asyncio
try:
    from youtube_transcript_api import YouTubeTranscriptApi  # type: ignore
    from youtube_transcript_api._errors import (  # type: ignore
        NoTranscriptFound,
        TranscriptsDisabled,
        CouldNotRetrieveTranscript,
        VideoUnavailable,
    )
    _yt_transcript_available = True
except Exception:
    _yt_transcript_available = False
    logger.warning("youtube-transcript-api is not available. Transcript extraction disabled.")

def _expand_language_preferences(prefer_langs: Optional[List[str]]) -> List[str]:
    """
    将语言偏好扩展成常见变体，并保持顺序与去重。
    例如: ['zh','en'] -> ['zh-Hans','zh-Hant','zh','zh-CN','zh-TW','en']
    """
    prefer_langs = prefer_langs or ["zh", "en"]
    expanded: List[str] = []
    for code in prefer_langs:
        c = (code or "").lower()
        variants: List[str]
        if c == "zh":
            variants = ["zh-Hans", "zh-Hant", "zh", "zh-CN", "zh-TW"]
        elif c == "en":
            variants = ["en"]
        else:
            variants = [code]
        for v in variants:
            if v not in expanded:
                expanded.append(v)
    return expanded

async def fetch_youtube_transcript(
    video_id: str,
    prefer_langs: Optional[List[str]] = None,
    fallback_any_language: bool = True,
) -> Dict[str, Any]:
    """
    仅通过字幕将视频转换为纯文本（不做语音转写）。
    策略：人工字幕优先，其次自动字幕；语言按 prefer_langs 顺序选择；若允许则回退到任意可用字幕。
    返回:
      { "text": str, "transcript_type": "human"|"auto", "lang": str, "orig_len": int }
    """
    if not _yt_transcript_available:
        raise ContentExtractionError(
            "youtube-transcript-api not installed. Please `pip install youtube-transcript-api`."
        )
    try:
        languages = _expand_language_preferences(prefer_langs)
        # 可选 cookies（用于通过登录态提升字幕命中率）
        cookies_path = os.getenv("YOUTUBE_COOKIES_FILE") or ""
        if cookies_path and not os.path.isfile(cookies_path):
            logger.warning(f"[youtube] cookies file not found: {cookies_path} (ignored)")
            cookies_path = ""
        if cookies_path:
            logger.info(f"[youtube] using cookies file for transcripts: {cookies_path}")

        transcripts = await asyncio.to_thread(
            YouTubeTranscriptApi.list_transcripts,
            video_id,
            None,
            cookies_path or None,
        )

        selected = None
        transcript_type = ""
        # 先尝试人工字幕
        try:
            selected = transcripts.find_transcript(languages)
            transcript_type = "human"
        except Exception:
            # 再尝试自动字幕
            try:
                selected = transcripts.find_generated_transcript(languages)
                transcript_type = "auto"
            except Exception:
                if fallback_any_language:
                    # 任意人工字幕
                    try:
                        manual = [t for t in transcripts if not getattr(t, "is_generated", False)]
                        if manual:
                            selected = manual[0]
                            transcript_type = "human"
                        else:
                            auto = [t for t in transcripts if getattr(t, "is_generated", False)]
                            if auto:
                                selected = auto[0]
                                transcript_type = "auto"
                    except Exception:
                        selected = None
        if not selected:
            raise ContentExtractionError("该视频无可用字幕，暂不支持（后续将支持音频转写）")

        # 主路径：优先使用选中的 Transcript.fetch()
        merged = ""
        try:
            items = await asyncio.to_thread(selected.fetch)
            lines: List[str] = []
            for it in items or []:
                txt = (it.get("text") or "").strip()
                if txt:
                    lines.append(txt)
            merged = "\n".join(lines).strip()
        except Exception:
            merged = ""

        # 兜底：遇到解析异常或空结果时，改走 get_transcript（更通用的 API）
        if not merged:
            try:
                raw_items = await asyncio.to_thread(
                    YouTubeTranscriptApi.get_transcript,
                    video_id,
                    languages=languages,
                    cookies=cookies_path or None,
                )
                lines: List[str] = []
                for it in raw_items or []:
                    txt = (it.get("text") or "").strip()
                    if txt:
                        lines.append(txt)
                merged = "\n".join(lines).strip()
            except Exception:
                merged = ""

        if not merged:
            # 最后兜底：使用 yt-dlp 下载字幕文件（无需 ffmpeg）
            try:
                merged = _download_captions_with_ytdlp(video_id, languages, cookies_path or None)
            except Exception as e:
                logger.warning(f"[youtube] yt-dlp caption fallback failed: {e}")
                merged = ""

        if not merged:
            raise ContentExtractionError("字幕拉取为空或被限制（稍后重试或更换视频试试）")
        return {
            "text": merged,
            "transcript_type": transcript_type or ("auto" if getattr(selected, "is_generated", False) else "human"),
            "lang": getattr(selected, "language_code", "") or "",
            "orig_len": len(merged),
        }
    except (NoTranscriptFound, TranscriptsDisabled):
        raise ContentExtractionError("该视频无字幕或字幕被禁用，暂不支持（后续将支持音频转写）")
    except (CouldNotRetrieveTranscript, VideoUnavailable) as e:
        raise ContentExtractionError(f"字幕获取失败：{str(e)}")
    except ContentExtractionError:
        raise
    except Exception as e:
        logger.exception(f"Unexpected error fetching transcript for video_id={video_id}")
        raise ContentExtractionError(f"字幕提取异常：{str(e)}")

# Step 4: 字幕清洗与规范化（独立函数，供上层在 Step 3 之后调用）
def clean_transcript_text(raw_text: str) -> str:
    """
    轻量清洗：
    - 去除常见方括号提示（[Music]、[Applause]、[Laughter] 等）
    - 压缩多余空格
    - 折叠多空行
    """
    text = raw_text or ""
    text = re.sub(r"\s*\[(?:music|applause|laughter|silence|ambient noise|.+?)\]\s*", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def normalize_transcript_text(text: str) -> str:
    """
    进一步规范化：
    - 去零宽/不可见字符（\u200b, \ufeff 等）
    - 连续重复行去重
    - 统一空白与空行
    """
    s = (text or "").replace("\u200b", "").replace("\ufeff", "")
    lines = [ln.strip() for ln in s.splitlines()]
    normalized: List[str] = []
    prev = None
    for ln in lines:
        if not ln:
            if normalized and normalized[-1] != "":
                normalized.append("")
            continue
        if ln == prev:
            continue
        normalized.append(ln)
        prev = ln
    out = "\n".join(normalized)
    out = re.sub(r"\n{3,}", "\n\n", out)
    out = re.sub(r"[ \t]{2,}", " ", out)
    return out.strip()

def clean_and_normalize_transcript(text: str) -> str:
    """
    组合式处理：先基础清洗，再规范化，得到适合送入模型的文本。
    """
    return normalize_transcript_text(clean_transcript_text(text or ""))

def _parse_vtt_to_text(vtt_content: str) -> str:
    """
    轻量解析 .vtt：去头、时间轴与序号，仅保留正文行。
    """
    body_lines = []
    for raw in (vtt_content or "").splitlines():
        s = raw.strip()
        if not s:
            continue
        if s.upper().startswith("WEBVTT"):
            continue
        if "-->" in s:  # 时间轴
            continue
        if s.isdigit():
            continue
        body_lines.append(s)
    return clean_and_normalize_transcript("\n".join(body_lines))

def _download_captions_with_ytdlp(video_id: str, languages: List[str], cookies_path: Optional[str]) -> str:
    """
    使用 yt-dlp 下载字幕（优先 --write-sub，失败试 --write-auto-sub），返回解析后的纯文本。
    """
    try:
        import yt_dlp  # noqa: F401
    except Exception:
        raise ContentExtractionError("yt-dlp 未安装，无法使用字幕兜底（pip install yt-dlp）")

    url = f"https://www.youtube.com/watch?v={video_id}"
    lang_csv = ",".join(languages or ["zh-Hans", "zh", "en"])
    with tempfile.TemporaryDirectory(prefix="ytvtt_") as td:
        out_tpl = str(Path(td) / "%(id)s.%(ext)s")
        base = [
            "yt-dlp",
            "--skip-download",
            "--sub-format", "vtt",
            "--sub-lang", lang_csv,
            "-o", out_tpl,
            url,
        ]
        if cookies_path:
            base += ["--cookies", cookies_path]
        # 先人工
        cmd1 = base + ["--write-sub"]
        r1 = subprocess.run(cmd1, capture_output=True, text=True)
        vtt_files = list(Path(td).glob(f"{video_id}*.vtt"))
        if not vtt_files:
            # 再自动
            cmd2 = base + ["--write-auto-sub"]
            r2 = subprocess.run(cmd2, capture_output=True, text=True)
            vtt_files = list(Path(td).glob(f"{video_id}*.vtt"))
            if not vtt_files:
                raise ContentExtractionError("yt-dlp 无法获取字幕（人工/自动均失败）")
        # 选择与优先语言最匹配的文件
        target = None
        for code in (languages or []):
            c = code.lower()
            for f in vtt_files:
                if c in f.name.lower():
                    target = f
                    break
            if target:
                break
        if target is None:
            target = vtt_files[0]
        content = target.read_text(encoding="utf-8", errors="ignore")
        parsed = _parse_vtt_to_text(content)
        if not parsed.strip():
            raise ContentExtractionError("已下载字幕但内容为空")
        return parsed

# Step 5: 长度策略（v1：字符上限，支持截断/单次摘要）
from app.services.llms import rewriting_client  # 延迟导入，避免不必要的依赖加载

def estimate_source_length_chars(text: str) -> int:
    """
    粗略长度估算，按字符数返回，用于 v1 的简单上限控制。
    """
    return len(text or "")

async def apply_length_policy(
    text: str,
    mode: str = "truncate",
    max_chars: int = 12000,
    llm_type: Optional[str] = None,
    summarize_instruction: Optional[str] = None,
) -> Dict[str, Any]:
    """
    对源文本应用长度策略：
    - mode == "truncate": 超过 max_chars 直接截断到上限
    - mode == "single_summarize": 超过上限时先用同一 LLM 做一次摘要压缩，再作为结果返回

    返回:
      {
        "text": str,          # 处理后的文本
        "mode": str,          # "original" | "truncate" | "single_summarize"
        "orig_len": int,      # 原始长度（字符）
        "final_len": int,     # 最终长度（字符）
        "truncated": bool     # 是否发生截断（仅 truncate 模式）
      }
    """
    src = text or ""
    orig_len = estimate_source_length_chars(src)
    if orig_len <= max_chars:
        return {
            "text": src,
            "mode": "original",
            "orig_len": orig_len,
            "final_len": orig_len,
            "truncated": False,
        }

    if mode == "truncate":
        out = src[:max_chars]
        return {
            "text": out,
            "mode": "truncate",
            "orig_len": orig_len,
            "final_len": len(out),
            "truncated": True,
        }

    if mode == "single_summarize":
        if not llm_type:
            raise ContentExtractionError("single_summarize requires llm_type")
        # 默认的摘要指令（可传入自定义 instruction 覆盖）
        instruction = summarize_instruction or (
            f"请将下面的内容压缩为精炼摘要，保留关键信息、实体与结构，长度不超过约 {max_chars} 个字符。"
        )
        try:
            logger.info("[length_policy] single_summarize start | llm_type={} | orig_len={} | max_chars={}", llm_type, orig_len, max_chars)
            summarized = await rewriting_client.get_rewriting_result(
                llm_type=llm_type,
                instruction=instruction,
                source=src,
            )
            final_text = summarized or ""
            return {
                "text": final_text,
                "mode": "single_summarize",
                "orig_len": orig_len,
                "final_len": len(final_text),
                "truncated": False,
            }
        except Exception as e:
            logger.exception("[length_policy] single_summarize failed | reason={}", e)
            raise ContentExtractionError(f"Summarization failed: {str(e)}")

    # 未知模式：回退截断
    out = src[:max_chars]
    return {
        "text": out,
        "mode": "truncate",
        "orig_len": orig_len,
        "final_len": len(out),
        "truncated": True,
    }

# Step 6: 编排（串联 1-5 步），面向路由的统一入口
async def ingest_youtube_url_v1(
    url: str,
    prefer_langs: Optional[List[str]] = None,
    fallback_any_language: bool = True,
    length_mode: str = "truncate",          # "truncate" | "single_summarize"
    max_chars: int = 12000,
    llm_type_for_summarize: Optional[str] = None,
    summarize_instruction: Optional[str] = None,
) -> Dict[str, Any]:
    """
    v1 组合流程：校验 -> 探测 -> 拉字幕 -> 清洗规范化 -> 长度策略
    返回:
      {
        "text": str,   # 清洗/限长之后的文本
        "meta": {
          "videoId": str,
          "title": str,
          "duration": int,
          "transcript_type": "human"|"auto",
          "lang": str,
          "length_mode": str,
          "orig_len": int,
          "final_len": int,
          "truncated": bool
        }
      }
    """
    # Step 1
    vid = validate_and_get_video_id(url)
    logger.info("[youtube] step1 ok | video_id={}", vid)
    # Step 2
    basic = probe_youtube_basic_info(url)
    logger.info("[youtube] step2 ok | video_id={} | title={} | duration={} | availability={}", basic.get("videoId"), basic.get("title"), basic.get("duration"), basic.get("availability"))
    # Step 3
    tr = await fetch_youtube_transcript(
        video_id=basic.get("videoId") or vid,
        prefer_langs=prefer_langs,
        fallback_any_language=fallback_any_language,
    )
    # Step 4
    cleaned = clean_and_normalize_transcript(tr.get("text") or "")
    # Step 5
    applied = await apply_length_policy(
        text=cleaned,
        mode=length_mode,
        max_chars=max_chars,
        llm_type=llm_type_for_summarize,
        summarize_instruction=summarize_instruction,
    )
    logger.info("[youtube] step3-5 ok | video_id={} | lang={} | type={} | final_len={} | mode={}", basic.get("videoId"), tr.get("lang"), tr.get("transcript_type"), applied.get("final_len"), applied.get("mode"))
    return {
        "text": applied["text"],
        "meta": {
            "videoId": basic.get("videoId") or vid,
            "title": basic.get("title") or "",
            "duration": basic.get("duration") or 0,
            "transcript_type": tr.get("transcript_type") or "",
            "lang": tr.get("lang") or "",
            "length_mode": applied.get("mode") or "",
            "orig_len": applied.get("orig_len") or 0,
            "final_len": applied.get("final_len") or 0,
            "truncated": bool(applied.get("truncated")),
        },
    }
